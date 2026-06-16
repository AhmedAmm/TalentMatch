"""
cv_generator.py
===============
Generates a tailored Canadian-style CV (DOCX) — pure python-docx, no Node.js.

Pipeline:
  1. MongoDB     — fetch employee by employee_id + job by job_id
  2. KG scoring  — rank skills by mastery × job-relevance  (kg_score + profil_score)
  3. PDF reader  — extract contact info from reference CV  (pdfplumber / pypdf)
  4. LLM         — generate summary, bullets, structured JSON
  5. python-docx — render polished Canadian Letter DOCX

Usage (CLI):
    python cv_generator.py \
        --employee-id  ahmed.ammar@talan.com \
        --job-id       proj_smartstaff_ing_nieur_ml_nlp_1 \
        --language     en \
        --cv           /path/to/reference_cv.pdf \
        --output       ahmed_cv_ml.docx

Usage (import):
    from cv_generator import generate_cv
    path = generate_cv(
        employee_id="ahmed.ammar@talan.com",
        job_id="proj_smartstaff_ing_nieur_ml_nlp_1",
        language="en",          # or "fr"
        cv_pdf_path="ref.pdf",  # optional
        output_path="out.docx",
    )

Install:
    pip install python-docx pymongo python-dotenv pdfplumber pypdf
"""

from __future__ import annotations

import argparse, json, os, re
from pathlib import Path
from textwrap import dedent
from typing import Optional

from dotenv import load_dotenv

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from services.llm import ask_llm

load_dotenv()

NAVY = RGBColor(0x1F, 0x4E, 0x79)
BLUE = RGBColor(0x2E, 0x75, 0xB6)
GRAY = RGBColor(0x40, 0x40, 0x40)

SECTION_LABELS = {
    "en": {
        "summary":        "Professional Summary",
        "skills":         "Core Technical Skills",
        "experience":     "Professional Experience",
        "education":      "Education",
        "certifications": "Certifications",
    },
    "fr": {
        "summary":        "Profil Professionnel",
        "skills":         "Compétences Techniques",
        "experience":     "Expérience Professionnelle",
        "education":      "Formation",
        "certifications": "Certifications",
    },
}

# ── DOCX helpers ──────────────────────────────────────────────────────────────

def _page_letter(doc):
    s = doc.sections[0]
    s.page_width = Inches(8.5); s.page_height = Inches(11)
    s.top_margin = s.bottom_margin = Inches(0.75)
    s.left_margin = s.right_margin = Inches(0.85)

def _rule(doc, color="2E75B6"):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), color)
    bdr.append(bot); pPr.append(bdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)

def _heading(doc, text):
    p   = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True; run.font.size = Pt(12)
    run.font.color.rgb = NAVY; run.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(2)
    _rule(doc)

def _para(doc, text, *, size=10, bold=False, italic=False,
          color=None, align="left", before=2, after=2):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size); run.bold = bold; run.italic = italic
    run.font.name = "Calibri"; run.font.color.rgb = color or GRAY
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def _bullet(doc, text):
    p   = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10); run.font.name = "Calibri"
    run.font.color.rgb = GRAY
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Inches(0.2)

def _mixed(doc, parts, before=6, after=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    for pt in parts:
        r = p.add_run(pt["text"])
        r.font.name = "Calibri"; r.font.size = Pt(pt.get("size", 10))
        r.bold = pt.get("bold", False); r.italic = pt.get("italic", False)
        r.font.color.rgb = pt.get("color", GRAY)

# ── 1. PDF reader ─────────────────────────────────────────────────────────────

def extract_pdf_text(path):
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            return "\n\n".join(pg.extract_text() or "" for pg in pdf.pages)
    except ImportError:
        pass
    try:
        from pypdf import PdfReader
        return "\n\n".join(pg.extract_text() or "" for pg in PdfReader(path).pages)
    except ImportError:
        print("[WARNING] Install pdfplumber or pypdf to use reference PDF.")
        return ""

def parse_contact(pdf_text, email):
    c = {"email": email}
    for line in pdf_text.split("\n"):
        line = line.strip()
        if line and "@" not in line and not re.search(r"\d{3}", line):
            c.setdefault("name", line); break
    m = re.search(r"(\+?[\d\s\-\(\)]{7,18})", pdf_text)
    if m: c["phone"] = m.group(1).strip()
    m = re.search(r"linkedin\.com/in/[\w\-]+", pdf_text, re.I)
    if m: c["linkedin"] = "https://" + m.group(0)
    m = re.search(r"\b([A-ZÀ-Ü][a-zà-ü]+(?:\s[A-ZÀ-Ü][a-zà-ü]+)*),\s*([A-Z]{2})\b", pdf_text)
    if m: c["location"] = f"{m.group(1)}, {m.group(2)}"
    return c

# ── 2. Fetch from MongoDB ──────────────────────────────────────────────────────

def build_job_description(job: dict) -> str:
    """
    Convert a job document into a plain-text description
    to feed the KG scorer and LLM prompt.
    """
    lines = [
        f"Title: {job.get('title', '')}",
        f"Seniority: {job.get('seniority', '')}",
        f"Type: {job.get('type', '')}",
        f"Description: {job.get('description', '')}",
        "",
        "Required stack:",
    ]
    for item in job.get("required_stack", []):
        lines.append(f"  - {item.get('skill', '')} ({item.get('level', '')})")

    lines.append("")
    lines.append("Responsibilities:")
    for r in job.get("responsibilities", []):
        lines.append(f"  - {r}")

    return "\n".join(lines)

# ── 3. KG scoring ─────────────────────────────────────────────────────────────

def relevant_skills(employee_doc: dict, job: dict, top_n: int = 20) -> list:
    """Rank employee skills by relevance to the job — standalone, no external state."""
    required = {
        (s["skill"] if isinstance(s, dict) else s).strip().lower()
        for s in (job.get("required_stack") or job.get("required_skills") or [])
        if s
    }
    ranked = []
    for skill in employee_doc.get("skills", []):
        name = (skill if isinstance(skill, str) else skill.get("name", "")).strip()
        if not name:
            continue
        is_match = name.lower() in required
        ranked.append({
            "name":        name,
            "coefficient": 1.0,
            "origin":      "direct",
            "relevance":   1.0 if is_match else 0.3,
        })
    ranked.sort(key=lambda x: x["relevance"], reverse=True)
    return ranked[:top_n]

# ── 4. LLM ────────────────────────────────────────────────────────────────────

def llm_cv_json(doc: dict, job: dict, skills: list, language: str, pdf_text: str) -> dict:
    lang = "English (Canadian)" if language == "en" else "French (Canadian)"

    projects = []
    for p in doc.get("projects", [])[:6]:
        notes = [t.get("description", "") for t in p.get("tasks", []) if isinstance(t, dict) and t.get("description")][:3]
        projects.append({
            "title":        p.get("name", p.get("project_id", "?")),
            "client":       p.get("client", ""),
            "role":         p.get("role", ""),
            "start":        p.get("start_date", ""),
            "end":          p.get("end_date") or "Present",
            "technologies": p.get("technologies", [])[:8],
            "task_notes":   notes,
        })

    required_stack = ", ".join(
        f"{s.get('skill')} ({s.get('level')})"
        for s in job.get("required_stack", [])
    )

    prompt = f"""You are an expert technical CV writer specialising in Canadian-style CVs.

TARGET JOB:
  Title       : {job.get('title', '')}
  Seniority   : {job.get('seniority', '')}
  Description : {job.get('description', '')}
  Required    : {required_stack}
  Responsibilities:
{chr(10).join('    - ' + r for r in job.get('responsibilities', []))}

ENGINEER SKILLS (job fit × mastery order):
{json.dumps([s["name"] for s in skills], indent=2)}

PROJECT HISTORY:
{json.dumps(projects, indent=2)}

REFERENCE CV TEXT:
{pdf_text[:3000] if pdf_text else "(not provided)"}

CURRENT ROLE: {doc.get("current_role", "")}
EDUCATION: {json.dumps(doc.get("education", []))}
CERTIFICATIONS: {json.dumps(doc.get("certifications", []))}

---
TASK: Write a tailored Canadian-style CV in {lang} for the engineer above.
Return ONLY valid JSON — no markdown fences, no preamble:
{{
  "summary": "4-5 sentence summary tailored to the job",
  "skills": ["skill1", "skill2", "..."],
  "experience": [
    {{
      "title": "Job Title", "company": "Company", "location": "City, XX",
      "start": "YYYY-MM", "end": "YYYY-MM or Present",
      "bullets": ["Action verb + achievement + metric", "..."]
    }}
  ],
  "education": [{{"degree":"","field":"","school":"","year":"YYYY"}}],
  "certifications": [{{"name":"","issuer":"","date":"YYYY"}}]
}}
RULES: bullets start with past-tense action verb, quantify where possible,
include ALL provided skills, do NOT invent data, output in {lang}."""

    raw = ask_llm(prompt, json_mode=True)
    raw = re.sub(r"```json\s*|```\s*", "", raw).strip()
    s = raw.find("{"); e = raw.rfind("}") + 1
    if s == -1 or e == 0:
        raise ValueError(f"LLM returned no JSON:\n{raw[:400]}")
    return json.loads(raw[s:e])

# ── 5. Render DOCX ────────────────────────────────────────────────────────────

def render_docx(cv: dict, contact: dict, labels: dict, out_path: str):
    doc = Document()
    _page_letter(doc)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # Name
    p   = doc.add_paragraph()
    run = p.add_run(contact.get("name", ""))
    run.bold = True; run.font.size = Pt(24)
    run.font.name = "Calibri"; run.font.color.rgb = NAVY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)

    if contact.get("current_role"):
        _para(doc, contact["current_role"], size=12, color=BLUE, align="center", before=0, after=2)

    parts = [contact["email"]]
    for k in ("phone", "location", "linkedin"):
        if contact.get(k): parts.append(contact[k])
    _para(doc, "  |  ".join(parts), size=9, align="center", before=0, after=4)
    _rule(doc, "1F4E79")

    # Summary
    _heading(doc, labels["summary"])
    _para(doc, cv.get("summary", ""), before=4, after=2)

    # Skills — 4 per row
    _heading(doc, labels["skills"])
    sk = cv.get("skills", [])
    for i in range(0, len(sk), 4):
        _para(doc, "   •   ".join(sk[i:i+4]), before=1, after=1)

    # Experience
    _heading(doc, labels["experience"])
    for exp in cv.get("experience", []):
        _mixed(doc, [
            {"text": (exp.get("title") or "") + " — ", "bold": True, "size": 11, "color": NAVY},
            {"text": exp.get("company") or "",          "bold": True, "size": 11, "color": GRAY},
        ], before=8, after=1)
        date = " – ".join(filter(None, [exp.get("start"), exp.get("end")]))
        meta = "   |   ".join(filter(None, [date, exp.get("location", "")]))
        _para(doc, meta, size=9, italic=True, color=BLUE, before=0, after=2)
        for b in exp.get("bullets", []):
            _bullet(doc, b)

    # Education
    if cv.get("education"):
        _heading(doc, labels["education"])
        for e in cv["education"]:
            deg = " in ".join(filter(None, [e.get("degree"), e.get("field")]))
            _para(doc, deg, size=11, bold=True, color=NAVY, before=6, after=1)
            school = ", ".join(filter(None, [e.get("school"), str(e.get("year", ""))]))
            _para(doc, school, size=10, italic=True, before=0, after=2)

    # Certifications
    if cv.get("certifications"):
        _heading(doc, labels["certifications"])
        for c in cv["certifications"]:
            line = "  —  ".join(filter(None, [c.get("name"), c.get("issuer"), str(c.get("date", ""))]))
            _para(doc, "• " + line, before=2, after=2)

    doc.save(out_path)
    print(f"[5/5] DOCX written → {out_path}")

# ── 6. Public API ─────────────────────────────────────────────────────────────

def generate_cv(
    employee_id: str,
    job_id: Optional[str],
    language: str = "en",
    cv_pdf_path: Optional[str] = None,
    output_path: Optional[str] = None,
    employee_doc: Optional[dict] = None,
    job_doc: Optional[dict] = None,
) -> str:
    """
    Full pipeline: MongoDB (employee + job) + KG scoring + LLM → Canadian DOCX.

    Parameters
    ----------
    employee_id  : employee MongoDB _id (email)
    job_id       : job MongoDB _id (e.g. proj_smartstaff_ing_nieur_ml_nlp_1)
    language     : "en" (default) or "fr"
    cv_pdf_path  : reference PDF CV path (optional — used for contact info)
    output_path  : output .docx path (default: <name>_<job_id>.docx)
    employee_doc : pre-fetched employee dict (skips DB fetch when provided)
    job_doc      : pre-fetched job dict (skips DB fetch when provided)

    Returns
    -------
    str — absolute path to the generated DOCX file
    """
    language = language.lower().strip()
    if language not in ("en", "fr"):
        raise ValueError("language must be 'en' or 'fr'.")

    labels = SECTION_LABELS[language]

    # ── Step 1: employee doc ───────────────────────────────────────────────────
    doc = employee_doc
    if doc is None:
        raise ValueError(
            f"No employee document provided for '{employee_id}'. "
            "Pass employee_doc= from an async context."
        )
    print(f"\n[1/5] Employee '{employee_id}' ready.")

    # ── Step 2: job doc & description ─────────────────────────────────────────
    job = job_doc
    if job is None and job_id:
        raise ValueError(
            f"No job document provided for '{job_id}'. "
            "Pass job_doc= from an async context."
        )
    if job is None:
        job = {}

    job_description = build_job_description(job)
    print(f"[2/5] Job '{job_id}': {job.get('title', '(no title)')}")

    # ── Default output path ────────────────────────────────────────────────────
    if output_path is None:
        name_slug = employee_id.split("@")[0].replace(".", "_")
        output_path = f"{name_slug}_{job_id or 'general'}.docx"

    # ── Step 3: KG scoring ─────────────────────────────────────────────────────
    print("[3/5] KG scoring — ranking skills …")
    skills = relevant_skills(doc, job) if job else []
    print(f"       Top: {', '.join(s['name'] for s in skills[:6])} …")

    # ── Step 4: PDF contact info ───────────────────────────────────────────────
    print("[4/5] PDF — reading reference CV …")
    pdf_text = ""
    if cv_pdf_path and Path(cv_pdf_path).exists():
        pdf_text = extract_pdf_text(cv_pdf_path)
        contact  = parse_contact(pdf_text, employee_id)
    else:
        if cv_pdf_path:
            print(f"       [WARNING] PDF not found: {cv_pdf_path}")
        contact = {"email": employee_id}
    contact.setdefault("name",         doc.get("name", employee_id))
    contact.setdefault("current_role", doc.get("current_role", ""))

    # ── Step 5: LLM generate ──────────────────────────────────────────────────
    print("[5/5] LLM — generating tailored CV content …")
    cv_data = llm_cv_json(doc, job, skills, language, pdf_text)

    # Back-fill from MongoDB if LLM omitted sections
    if not cv_data.get("education") and doc.get("education"):
        cv_data["education"] = [
            {"degree": e.get("degree", ""), "field": e.get("field", ""),
             "school": e.get("school", ""), "year":  str(e.get("year", ""))}
            for e in doc["education"]
        ]
    if not cv_data.get("certifications") and doc.get("certifications"):
        cv_data["certifications"] = [
            {"name": c.get("name", ""), "issuer": c.get("issuer", ""),
             "date": str(c.get("date", ""))}
            for c in doc["certifications"]
        ]

    # ── Step 6: render DOCX ───────────────────────────────────────────────────
    render_docx(cv_data, contact, labels, output_path)

    abs_path = str(Path(output_path).resolve())
    print(f"\n✅  Done — {abs_path}\n")
    return abs_path

# ── CLI ───────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    ap = argparse.ArgumentParser(
        description="Generate a tailored Canadian-style CV (DOCX) from employee + job IDs.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=dedent("""
        Examples:
          python cv_generator.py \\
              --employee-id ahmed.ammar@talan.com \\
              --job-id      proj_smartstaff_ing_nieur_ml_nlp_1 \\
              --language    en --cv ref.pdf

          python cv_generator.py \\
              --employee-id amir.bensalah@email.com \\
              --job-id      proj_smartstaff_tech_lead_architecte_ia_0 \\
              --language    fr
        """),
    )
    ap.add_argument("--employee-id", required=True, help="Employee email (_id in employees collection)")
    ap.add_argument("--job-id",      required=True, help="Job _id (e.g. proj_smartstaff_ing_nieur_ml_nlp_1)")
    ap.add_argument("--language",    default="en",  help="'en' or 'fr'  (default: en)")
    ap.add_argument("--cv",          default=None,  help="Reference PDF CV path (optional)")
    ap.add_argument("--output",      default=None,  help="Output .docx path (default: <name>_<job_id>.docx)")
    args = ap.parse_args()

    generate_cv(
        employee_id=args.employee_id,
        job_id=args.job_id,
        language=args.language,
        cv_pdf_path=args.cv,
        output_path=args.output,
    )